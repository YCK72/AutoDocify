import json
from pathlib import Path
from docx import Document

def load_metadata(json_file):
    with open(json_file, 'r', encoding='utf-8') as f:
        return json.load(f)

def export_to_markdown(metadata, md_file_path):
    lines = [f"# Documentation for `{Path(metadata['file_path']).name}`\n"]

    functions = metadata.get("ast_data", {}).get("functions", [])
    if functions:
        lines.append("## Functions")
        for f in functions:
            lines.append(f"### `{f['name']}`\n")
            lines.append(f"```python\n{f['docstring'].strip('\"')}\n```\n")

    classes = metadata.get("ast_data", {}).get("classes", [])
    if classes:
        lines.append("## Classes")
        for c in classes:
            lines.append(f"### `{c['name']}`\n")
            lines.append(f"```python\n{c['docstring'].strip('\"')}\n```\n")
            for m in c.get("methods", []):
                lines.append(f"#### Method `{m['name']}`\n")
                lines.append(f"```python\n{m['docstring'].strip('\"')}\n```\n")

    Path(md_file_path).parent.mkdir(parents=True, exist_ok=True)
    with open(md_file_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"✅ Markdown saved to: {md_file_path}")

def export_to_word(metadata, docx_path):
    doc = Document()
    doc.add_heading(f"Documentation for {Path(metadata['file_path']).name}", 0)

    functions = metadata.get("ast_data", {}).get("functions", [])
    if functions:
        doc.add_heading("Functions", level=1)
        for f in functions:
            doc.add_heading(f"{f['name']}", level=2)
            doc.add_paragraph(f['docstring'].strip('"'), style='Normal')

    classes = metadata.get("ast_data", {}).get("classes", [])
    if classes:
        doc.add_heading("Classes", level=1)
        for c in classes:
            doc.add_heading(c["name"], level=2)
            doc.add_paragraph(c["docstring"].strip('"'), style='Normal')
            for m in c.get("methods", []):
                doc.add_heading(f"Method: {m['name']}", level=3)
                doc.add_paragraph(m['docstring'].strip('"'), style='Normal')

    Path(docx_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"✅ Word document saved to: {docx_path}")
