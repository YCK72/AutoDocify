import json
from docx import Document
from pathlib import Path

def json_to_docx(json_path, docx_path):
    data = json.load(open(json_path, 'r', encoding='utf-8'))
    doc = Document()
    doc.add_heading('AutoDocify Documentation', level=1)

    # Adjust keys/structure to match your JSON schema
    for file_entry in data.get('files', []):
        doc.add_heading(file_entry['filename'], level=2)
        for section in file_entry.get('docs', []):
            doc.add_heading(section['type'], level=3)
            doc.add_paragraph(section['content'])

    doc.save(docx_path)

if __name__ == '__main__':
    import sys
    json_to_docx(sys.argv[1], sys.argv[2])